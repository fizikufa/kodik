
import io
import base64
import mimetypes
from pathlib import Path

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_TEXT_CHARS = 80_000

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt",
    ".png", ".jpg", ".jpeg", ".webp", ".gif",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}


async def extract_text(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Стр. {i+1}]\n{text}")
        return "\n\n".join(pages)

    elif ext == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                lines.append("\t".join(c.text.strip() for c in row.cells))
        return "\n".join(lines)

    elif ext == ".xlsx":
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            lines.append(f"=== Лист: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    lines.append("\t".join(cells))
        return "\n".join(lines)

    elif ext == ".xls":
        import xlrd
        wb = xlrd.open_workbook(file_contents=data)
        lines = []
        for sheet in wb.sheets():
            lines.append(f"=== Лист: {sheet.name} ===")
            for i in range(sheet.nrows):
                lines.append("\t".join(str(v) for v in sheet.row_values(i)))
        return "\n".join(lines)

    elif ext == ".csv":
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")

    elif ext == ".txt":
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")

    else:
        raise ValueError(f"Неподдерживаемый формат: {ext}")


def encode_image(filename: str, data: bytes) -> dict:
    mime = mimetypes.guess_type(filename)[0] or "image/jpeg"
    b64 = base64.b64encode(data).decode()
    return {"type": "image", "media_type": mime, "data": b64, "filename": filename}
